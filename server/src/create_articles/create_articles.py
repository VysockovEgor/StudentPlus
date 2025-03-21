
import openai
from openai import ChatCompletion  # Новый импорт для работы с ChatCompletion
from docx import Document
from dotenv import load_dotenv
import os

load_dotenv()
SECRET_KEY = os.getenv('OPENAI_API_KEY')
PROJECT_DOCS_DIRECTORY = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', '..', 'articles')

def load_project_documents():
    """
    Загружает содержимое всех .docx файлов из папки проекта и возвращает объединённый текст.
    """
    all_text = []
    if os.path.exists(PROJECT_DOCS_DIRECTORY):
        for filename in os.listdir(PROJECT_DOCS_DIRECTORY):
            if filename.endswith('.docx'):
                file_path = os.path.join(PROJECT_DOCS_DIRECTORY, filename)
                try:
                    doc = Document(file_path)
                    text = "\n".join([para.text for para in doc.paragraphs])
                    all_text.append(f"--- Document: {filename} ---\n{text}\n")
                except Exception as e:
                    print(f"Ошибка чтения {filename}: {e}")
    return "\n".join(all_text)


def generate_article(title, instructions, system_prompt=None, model="gpt-3.5-turbo"):
    """
    Генерирует статью с использованием OpenAI API.

    Параметры:
      - title: Заголовок статьи.
      - instructions: Инструкции по созданию статьи.
      - system_prompt: (Опционально) Системный промпт для генерации.
      - model: Модель для генерации (по умолчанию gpt-3.5-turbo).

    Возвращает:
      Сгенерированный текст статьи.
    """
    if system_prompt is None:
        system_prompt = (
            "You are a creative writer. Generate an engaging, informative article "
            "based on the given title, instructions, and project documents. "
            "Use the provided documents as a reference and ensure the article is coherent and well-structured."
        )

    # Загружаем все документы проекта
    project_docs_text = load_project_documents()

    # Формируем список сообщений для ChatCompletion API
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user",
         "content": f"Title: {title}\nInstructions: {instructions}\n\nProject Documents:\n{project_docs_text}"}
    ]


    response = ChatCompletion.create(
        model=model,
        messages=messages,
        temperature=0.7
    )

    generated_text = response.choices[0].message.content.strip()
    return generated_text


def save_article(title, content, directory=PROJECT_DOCS_DIRECTORY):
    """
    Сохраняет сгенерированную статью в новый .docx файл с именем, соответствующим заголовку.

    Параметры:
      - title: Заголовок статьи (будет использоваться как имя файла).
      - content: Содержимое статьи (текст).
      - directory: Папка для сохранения файла.
    """
    if not os.path.exists(directory):
        os.makedirs(directory)

    # Если заголовок не содержит расширение, добавляем .docx
    file_name = title if title.lower().endswith('.docx') else title + ".docx"
    file_path = os.path.join(directory, file_name)

    doc = Document()
    for line in content.split('\n'):
        doc.add_paragraph(line)
    doc.save(file_path)
    print(f"Статья сохранена в {file_path}")




